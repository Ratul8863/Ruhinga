# -*- coding: utf-8 -*-
"""Rebuild v06 from v05 with correct Step 3.3 documentation order."""
import docx
from docx.oxml.ns import qn
from docx.shared import Pt

PATH = "01_Report/Full_Research_Report_v06_Annual_Exposure_Reconstruction.docx"
SRC = "01_Report/Full_Research_Report_v05_Boundary_Stage_Frozen.docx"


def find(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise KeyError(text)


def insert_after(anchor, text, style="Normal"):
    p = anchor._parent.add_paragraph()
    p.style = style
    p.add_run(text)
    anchor._p.addnext(p._p)
    return p


def set_cell_font(cell, size=8):
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.size = Pt(size)


def add_table_after(doc, anchor, headers, rows, font_size=8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_font(table.rows[0].cells[i], font_size)
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
            set_cell_font(cells[i], font_size)
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    anchor._p.addnext(tbl)
    return table


def heading_style(doc, preferred="Heading 4"):
    names = {s.name for s in doc.styles}
    return preferred if preferred in names else "Heading 3"


def append_block(doc, anchor, items):
    """Append paragraphs/tables after anchor in forward order; return last paragraph anchor."""
    cur = anchor
    for item in items:
        kind = item[0]
        if kind == "p":
            _, style, text = item
            cur = insert_after(cur, text, style)
        elif kind == "table":
            _, caption, headers, rows = item
            cur = insert_after(cur, caption)
            add_table_after(doc, cur, headers, rows)
            # table sits after caption; keep caption as textual anchor for next inserts
            # Move next inserts after the table by finding the table element
            tbl = cur._p.getnext()
            while tbl is not None and tbl.tag != qn("w:tbl"):
                tbl = tbl.getnext()
            if tbl is None:
                raise RuntimeError("Table not found after caption")
            # Create a dummy empty paragraph after table to serve as next anchor
            spacer = insert_after(cur, "")
            # Currently: cur(caption), table?, spacer was insert_after(cur) so spacer is between caption and table!
            # Fix: re-locate. Better approach below.
            raise RuntimeError("Use append_block_v2")
    return cur


def chain_after(doc, start_anchor, items):
    """
    Insert items in forward order after start_anchor.
    Each new paragraph is inserted after the previous paragraph.
    Tables are inserted after their caption paragraph.
    Returns the last paragraph element used as anchor.
    """
    prev = start_anchor
    last_para = start_anchor
    for item in items:
        kind = item[0]
        if kind == "p":
            _, style, text = item
            p = insert_after(prev, text, style)
            # If prev was a caption and a table already follows prev, insert_after(prev)
            # places the new para between prev and table. To keep order para→table→next,
            # after adding a table we set prev to a marker after the table.
            prev = p
            last_para = p
        elif kind == "table":
            _, caption, headers, rows = item
            cap = insert_after(prev, caption)
            add_table_after(doc, cap, headers, rows)
            # Ensure next content goes after the table: create empty para after table
            tbl = cap._p.getnext()
            if tbl is None or tbl.tag != qn("w:tbl"):
                raise RuntimeError("Expected table after caption")
            marker = cap._parent.add_paragraph()
            marker.add_run("")
            tbl.addnext(marker._p)
            prev = marker
            last_para = marker
    return last_para


# Fresh copy from frozen v05
import shutil

shutil.copyfile(SRC, PATH)
doc = docx.Document(PATH)
h4 = heading_style(doc)

# Title page
for p in doc.paragraphs:
    if p.text.startswith("Version:"):
        p.text = "Version: v06"
    elif p.text.startswith("Date:"):
        p.text = "Date: 27 July 2026"
    elif p.text.startswith("Stage:"):
        p.text = "Stage: Step 3.3 Annual Camp Exposure Reconstruction — In Progress"
    elif p.text.startswith("□ Annual Camp Exposure Reconstruction"):
        p.text = "◐ Annual Camp Exposure Reconstruction — In Progress (Step 3.3)"

# Replace old 4.4.3 body
h443 = find(doc, "4.4.3 Annual Exposure Reconstruction (Future Work)")
h45 = find(doc, "4.5 Forest-Loss Analysis")
cur = h443._p.getnext()
while cur is not None and cur is not h45._p:
    nxt = cur.getnext()
    if cur.tag == qn("w:p"):
        cur.getparent().remove(cur)
    cur = nxt

h443.text = "4.4.3 Annual Exposure Reconstruction"

items_443 = [
    ("p", h4, "4.4.3.1 Reconstruction Processing AOI"),
    (
        "p",
        "Normal",
        "A provisional processing area was constructed for annual satellite-based "
        "camp-exposure reconstruction. The verified 2018, 2023 and 2024 A2-derived "
        "official camp layers were merged, producing 104 source features, and dissolved "
        "without a grouping field into a single multipart geometry. A 2,000 m buffer was "
        "then generated in WGS 84 / UTM Zone 46N (EPSG:32646) using round joins, 20 "
        "segments and a dissolved output.",
    ),
    (
        "p",
        "Normal",
        "The resulting layer, reconstruction_aoi_2016_2024_2km_final.gpkg, contained one "
        "valid multipart feature. Strict QGIS validity testing returned one valid feature, "
        "zero invalid features and zero geometry errors. The AOI covered approximately "
        "189.372 km².",
    ),
    (
        "p",
        "Normal",
        "This geometry is used only as a satellite-image processing and reconstruction AOI. "
        "It is not a final treatment, spillover or donor-zone boundary.",
    ),
    (
        "p",
        "Normal",
        "Output files — official_camp_snapshots_2018_2024_merged.gpkg; "
        "official_camp_snapshots_2018_2024_union.gpkg; "
        "reconstruction_aoi_2016_2024_2km_final.gpkg; "
        "reconstruction_aoi_2016_2024_2km_wgs84.zip. Path: "
        "04_Data/02_Camp_Exposure/Annual_Reconstruction/01_AOI/.",
    ),
    (
        "p",
        "Normal",
        "Earth Engine asset — projects/rohingya-forest-impact-2026/assets/"
        "reconstruction_aoi_2016_2024_2km.",
    ),
    ("p", h4, "4.4.3.2 2016 Imagery Source and Seasonal-Window Selection"),
    (
        "p",
        "Normal",
        "Landsat 8 Collection 2 Tier 1 Level 2 Surface Reflectance imagery was evaluated "
        "for the 2016 pre-influx settlement and land-surface baseline. Cloud, cirrus, "
        "dilated-cloud, cloud-shadow, snow, fill and radiometrically saturated pixels were "
        "removed using the QA_PIXEL and QA_RADSAT bands. Surface-reflectance bands were "
        "scaled using the Collection 2 Level 2 scale factor and offset.",
    ),
    (
        "p",
        "Normal",
        "Three candidate temporal windows were evaluated using scene count, scene-level "
        "cloud-cover statistics, AOI-based valid-pixel coverage and visual interpretation.",
    ),
    (
        "table",
        "Table. Candidate 2016 Landsat seasonal windows evaluated for the reconstruction AOI.",
        ["Candidate window", "Scene count", "Mean cloud cover", "Valid-pixel coverage"],
        [
            ["Jan–Apr 2016", "28", "22.74%", "99.738%"],
            ["Nov 2015–Apr 2016", "42", "16.96%", "99.738%"],
            ["Full year 2016", "69", "26.09%", "99.738%"],
        ],
    ),
    (
        "p",
        "Normal",
        "The period from 1 November 2015 to 30 April 2016 was selected. All candidate "
        "windows provided nearly identical valid-pixel coverage, but the selected window "
        "had the lowest mean scene cloud cover and retained a dry-season acquisition period "
        "without including any post-influx dates. Visual inspection found no obvious large "
        "cloud patches, major cloud-shadow contamination or large missing-data gaps.",
    ),
    (
        "p",
        "Normal",
        "Interpretation — The product is documented as a 2016 pre-influx settlement and "
        "land-surface baseline. It is not reported as a 2016 Rohingya camp footprint at "
        "this stage, because settlement classification and physical-footprint delineation "
        "remain pending.",
    ),
    ("p", h4, "4.4.3.3 Composite Construction and Export"),
    (
        "p",
        "Normal",
        "The selected 42 Landsat scenes were processed into a pixel-wise median Surface "
        "Reflectance composite. The composite was clipped to the reconstruction AOI and "
        "exported at 30 m resolution in EPSG:32646.",
    ),
    (
        "p",
        "Normal",
        "Six reflectance bands and four spectral indices were retained: blue, green, red, "
        "near-infrared, SWIR1, SWIR2, NDVI, NDBI, MNDWI and BSI. A valid-observation-count "
        "band was also included to record the number of cloud-free observations "
        "contributing to each pixel.",
    ),
    (
        "p",
        "Normal",
        "Output — L8_SR_2016_preinflux_windowB_composite_30m.tif. Path: "
        "04_Data/02_Camp_Exposure/Annual_Reconstruction/03_Annual_Composites/.",
    ),
    (
        "p",
        "Normal",
        "GEE scripts — 01_2016_Landsat_Window_Test; 02_2016_Landsat_Composite_Export "
        "(reproducibility archive).",
    ),
]

# Insert the block after h443, but physically before h45.
# chain_after inserts after h443; because h45 follows the removed body, new content
# will appear between h443 and h45 if we insert after h443.
chain_after(doc, h443, items_443)

# 4.12.3 after 4.12.2 interpretation
interp_4122 = None
for p in doc.paragraphs:
    if p.text.startswith("Interpretation — Shamlapur was the only 2018 site"):
        interp_4122 = p
        break
if interp_4122 is None:
    raise RuntimeError("4.12.2 interpretation not found")

items_4123 = [
    ("p", "Heading 3", "4.12.3 2016 Baseline Composite QA"),
    (
        "p",
        "Normal",
        "Purpose — This validation assessed geometry integrity of the reconstruction "
        "processing AOI and raster integrity of the exported 2016 pre-influx Landsat "
        "baseline composite.",
    ),
    (
        "p",
        "Normal",
        "The exported GeoTIFF loaded successfully in QGIS and contained 11 Float32 bands. "
        "The raster used EPSG:32646, a 30 × 30 m pixel size, dimensions of 610 columns × "
        "1,259 rows, GeoTIFF format and LZW compression.",
    ),
    (
        "p",
        "Normal",
        "Valid-observation counts ranged from 13 to 31 observations per pixel in Earth "
        "Engine, with a mean of approximately 26.02. QGIS approximate statistics returned "
        "a minimum of 14, mean of 26.01 and maximum of 31; the one-observation difference "
        "in the minimum was attributed to approximate raster-statistics sampling.",
    ),
    (
        "p",
        "Normal",
        "The reported QGIS valid-pixel percentage of approximately 27.33% reflected the "
        "rectangular raster extent surrounding a spatially disjoint multipart AOI. It did "
        "not represent cloud-free coverage within the AOI. AOI-based valid-pixel coverage "
        "calculated in Earth Engine was approximately 99.738%.",
    ),
    (
        "table",
        "Table. 2016 baseline composite and reconstruction-AOI QA summary.",
        ["QA item", "Result"],
        [
            ["AOI features", "1 multipart"],
            ["AOI area", "189.372 km²"],
            ["AOI valid / invalid / errors", "1 / 0 / 0"],
            ["Input Landsat scenes", "42"],
            ["Raster CRS", "EPSG:32646"],
            ["Pixel size", "30 × 30 m"],
            ["Raster dimensions", "610 × 1,259"],
            ["Band count", "11"],
            ["Data type", "Float32"],
            ["Valid observations", "13–31"],
            ["Mean valid observations", "26.02"],
            ["AOI valid-pixel coverage", "99.738%"],
        ],
    ),
    (
        "p",
        "Normal",
        "Interpretation — AOI geometry and the 2016 baseline composite passed initial "
        "integrity and metadata checks. Settlement classification, physical-footprint "
        "delineation and accuracy assessment remain pending; therefore, no 2016 camp-area "
        "or expansion result is reported.",
    ),
]
chain_after(doc, interp_4122, items_4123)

# Update 5.7 body
h57 = find(doc, "5.7 Empirical Work Remaining")
body57 = None
el = h57._p.getnext()
while el is not None:
    if el.tag == qn("w:p"):
        for p in doc.paragraphs:
            if p._p is el:
                body57 = p
                break
        break
    el = el.getnext()
if body57 is not None:
    body57.text = (
        "The 2018–2023 historical identifier crosswalk has been completed and validated, "
        "and Step 3.3 annual exposure reconstruction has been initiated. Remaining work "
        "includes settlement classification and physical-footprint delineation for 2016, "
        "reconstruction of other missing annual camp footprints (2017 and unavailable "
        "years), annual exposure reconstruction completion, Hansen outcome extraction, "
        "fragmentation analysis, SDID implementation, biomass/carbon estimation, "
        "machine-learning modelling and conservation-priority mapping."
    )

# 5.8 before chapter 6 — insert after 5.7 body using chain
h6 = find(doc, "6. Results")
# insert immediately before h6 by chaining after body57 if available, else after h57
anchor_58 = body57 if body57 is not None else h57
# But body57 may not be immediately before h6 (there may be nothing between). Use insert before h6 via reverse-safe method: insert_after body57.
chain_after(
    doc,
    anchor_58,
    [
        ("p", "Heading 2", "5.8 Annual Exposure Reconstruction Progress"),
        (
            "p",
            "Normal",
            "Step 3.3 was initiated through construction of a 189.372 km² "
            "satellite-processing AOI and preparation of a cloud-masked 2016 pre-influx "
            "Landsat baseline composite. The selected November 2015–April 2016 window "
            "contained 42 scenes and achieved approximately 99.738% valid-pixel coverage "
            "within the AOI. The exported 11-band, 30 m composite passed initial "
            "raster-integrity and metadata checks. Settlement classification and "
            "physical-footprint delineation remain pending; therefore, no 2016 camp-area "
            "or expansion result is reported at this stage.",
        ),
    ],
)

# Results note
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "6. Results":
        nxt = doc.paragraphs[i + 1]
        if nxt.text.strip().startswith("[To be completed") or "No empirical" in nxt.text:
            nxt.text = (
                "[No empirical camp-footprint, forest-loss, fragmentation or carbon "
                "results are reported at this stage. Section 4 documents methods and QA "
                "only; Section 5 documents research progress.]"
            )
        break

# Remove empty marker paragraphs created after tables (optional cleanup)
for p in list(doc.paragraphs):
    if p.text == "" and p.style.name == "Normal":
        # only remove markers that are truly empty and sit between table and next content
        # safer: leave them; they are harmless blank lines
        pass

doc.save(PATH)
print("Saved", PATH)
