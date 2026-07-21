# -*- coding: utf-8 -*-
"""Apply final QA corrections to the v02 working report only."""
import docx

PATH = "01_Report/Full_Research_Report_v02_Literature_and_Feasibility.docx"


def find(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise KeyError(text)


def insert_before(doc, anchor, text, style="Normal"):
    p = doc.add_paragraph()
    p.style = style
    p.add_run(text)
    anchor._p.addprevious(p._p)
    return p


def replace_between(doc, start, end, items):
    ps = list(doc.paragraphs)
    s = next(i for i, p in enumerate(ps) if p._p is start._p)
    e = next(i for i, p in enumerate(ps) if p._p is end._p)
    for p in ps[s + 1:e]:
        p._element.getparent().remove(p._element)
    for text, style in items:
        insert_before(doc, end, text, style)


doc = docx.Document(PATH)

# 1) Calibrate camp-boundary statements to metadata/source discovery rather than local vector verification.
for p in doc.paragraphs:
    if p.text.startswith("The final analytical grid will be constructed in WGS 84 / UTM Zone 46N"):
        p.text = (
            "The final analytical grid will be constructed in WGS 84 / UTM Zone 46N (EPSG:32646). "
            "Candidate official boundary sources and dated products from 2018, 2023 and 2024 have been "
            "identified and will be evaluated alongside satellite-derived annual exposure. Required vector "
            "files have not yet all been downloaded, inspected and spatially validated within the project workspace."
        )
    elif p.text.startswith("Annual camp exposure will combine official 2018, 2023 and 2024 boundary snapshots"):
        p.text = (
            "Annual camp exposure will combine candidate official boundary sources and dated products identified "
            "for 2018, 2023 and 2024 with satellite-derived 2016–2017 and other missing-year footprints. "
            "Before use, the required vector files must be downloaded, inspected and validated at file and "
            "geometry level within the project workspace. Population tables will be joined as exposure/context "
            "information. Raw official boundary files will remain unedited; processed exposure layers will be "
            "created separately."
        )
    elif p.text.startswith("Official camp-boundary snapshots are available from 2018 onward"):
        p.text = (
            "Official camp-boundary sources and dated products have been identified for 2018, 2023 and 2024. "
            "However, the required vector files have not yet all been downloaded, inspected and spatially "
            "validated within the project workspace. Their usability remains subject to file-level and "
            "geometry-level verification. Pre-influx 2016 and emergency-expansion 2017 footprints, plus any "
            "missing years, require satellite reconstruction. Population tables have been located but still "
            "require harmonisation and polygon linkage."
        )

# 2) Replace Appendix A using the current 6 objective / 6 RQ design.
appendix = find(doc, "Appendix A. Objective-Research Question Alignment Check")
replace_between(doc, appendix, None if False else appendix, [])  # no-op marker
# Appendix A is final section: delete all paragraphs after it and append the corrected mapping.
ps = list(doc.paragraphs)
idx = next(i for i, p in enumerate(ps) if p._p is appendix._p)
for p in ps[idx + 1:]:
    p._element.getparent().remove(p._element)
for text in [
    "Internal design check (for research-design verification; not required for public report text):",
    "Objective 1 — Annual forest outcomes → RQ1",
    "Objective 2 — Annual fragmentation and connectivity → RQ2",
    "Objective 3 — SDID causal attribution → RQ3",
    "Objective 4 — Biomass, carbon and CO₂-equivalent loss → RQ4",
    "Objective 5 — Machine learning and SHAP prediction → RQ5",
    "Objective 6 — Conservation and restoration priority → RQ6",
    "Alignment check: Correct. Every objective has a corresponding research question, and no research question falls outside the stated objectives.",
]:
    new = doc.add_paragraph(text)
    new.style = "Normal"

# 3) Replace References placeholder with the frozen review bibliography and dataset citations.
references = find(doc, "References")
appendices = find(doc, "Appendices")
reference_intro = (
    "Provisional reference list based on the frozen targeted literature review. Final journal formatting "
    "will be completed during manuscript preparation."
)
references_list = [
    "P001. Hassan, M. M., Smith, A. C., Walker, K., Rahman, M. K., & Southworth, J. (2018). Rohingya Refugee Crisis and Forest Cover Change in Teknaf, Bangladesh. Remote Sensing, 10(5), Article 689. https://doi.org/10.3390/rs10050689",
    "P002. Dampha, N. K., Salemi, C., & Polasky, S. (2022). Rohingya Refugee Camps and Forest Loss in Cox's Bazar, Bangladesh: An Inquiry Using Remote Sensing and Econometric Approaches. World Bank Policy Research Working Paper No. 9948. https://doi.org/10.1596/1813-9450-9948",
    "P003. Hassan, M. M., Duveneck, M., & Southworth, J. (2023). The Role of the Refugee Crises in Driving Forest Cover Change and Fragmentation in Teknaf, Bangladesh. Ecological Informatics. https://doi.org/10.1016/j.ecoinf.2022.101966",
    "P004. Rahaman, M., Morshed, M. M., & Bhadra, S. (2022). An integrated machine learning and remote sensing approach for monitoring forest degradation due to Rohingya refugee influx in Bangladesh. Remote Sensing Applications: Society and Environment. https://doi.org/10.1016/j.rsase.2022.100696",
    "P005. Mahmood, H., Saha, C., & Saha, S. (2025). Tracking forest recovery: Early biomass and carbon stock monitoring in the Rohingya Refugee camps, Cox's Bazar, Bangladesh. Environmental Challenges. https://doi.org/10.1016/j.envc.2024.101063",
    "P006. Mitra, J. R., Ahmed, T. T., & Czajkowski, K. (2025). Landscape fragmentation and population distribution in refugee camps: Evidence from the Rohingya refugee influx in Bangladesh. Population and Environment. https://doi.org/10.1007/s11111-025-00489-4",
    "P007. Arkhangelsky, D., Athey, S., Hirshberg, D. A., Imbens, G. W., & Wager, S. (2021). Synthetic Difference-in-Differences. American Economic Review. https://doi.org/10.1257/aer.20190159",
    "P008. Salemi, C. (2025). Rohingya Refugee Camps and Land Cover Change in Bangladesh [Working paper under review]. https://colettesalemi.ca/files/CB_LC_2025.pdf",
    "P009. Langer, S., Tiede, D., & Lüthje, F. (2015). Long-term Monitoring of the Environmental Impact of a Refugee Camp Based on Landsat Time Series: The Example of Deforestation and Reforestation During the Whole Lifespan of the Camp Lukole, Tanzania. GI_Forum. https://doi.org/10.1553/giscience2015s434",
    "P010. Hasan, M. E., Zhang, L., Dewan, A., Guo, H., & Mahmood, R. (2021). Spatiotemporal Pattern of Forest Degradation and Loss of Ecosystem Function Associated with Rohingya Influx: A Geospatial Approach. Land Degradation & Development, 32(13), 3666–3683. https://doi.org/10.1002/ldr.3821",
    "P011. Sarkar, S. K., Saroar, M., & Chakraborty, T. (2023b). Navigating nature’s toll: Assessing the ecological impact of the refugee crisis in Cox’s Bazar, Bangladesh. Heliyon. https://doi.org/10.1016/j.heliyon.2023.e18255",
    "P012. Sarkar, S. K., Saroar, M. M., & Chakraborty, T. (2023a). Cost of Ecosystem Service Value Due to Rohingya Refugee Influx in Bangladesh. Disaster Medicine and Public Health Preparedness. https://doi.org/10.1017/dmp.2022.125",
    "P013. Ahmed, F., Alam, S., Saha, O. R., & Rahman, A. (2024). The Rohingya refugee crisis in Bangladesh: Assessing the impact on land use patterns and land surface temperature using machine learning. Environmental Monitoring and Assessment. https://doi.org/10.1007/s10661-024-12701-3",
    "P014. Gromny, E., Jenerowicz-Sanikowska, M., Haarpaintner, J., Aleksandrowicz, S., Woźniak, E., Pesquer Mayos, L., Chułek, M., Sobczak-Szelc, K., Wawrzaszek, A., Sala, S., Espegren, A., Starczewski, D., & Pawlak, Z. (2024). Remote sensing insights into land cover dynamics and socio-economic drivers: The case of Mtendeli refugee camp, Tanzania (2016–2022). Remote Sensing Applications: Society and Environment. https://doi.org/10.1016/j.rsase.2024.101334",
    "P015. Mohiuddin, M., Hossain, M., Islam, M. Y., Nowreen, S., & Sultana, N. (2025). Assessing the Impact of Land Use Land Cover Change after Rohingya Forced Migration on Elephant Route in Cox's Bazar District, Bangladesh. Landscape and Ecological Engineering. https://doi.org/10.1007/s11355-025-00642-z",
    "P016. Chowdhury, F. I., Bhuiyan, R. H., Espelta, J. M., Resco de Dios, V., Dilshad, T., Haque, M. R., Aman, M. A. U., & Lloret, F. (2025). Land-use legacies and tree species richness affect short-term resilience in reforested areas of the world’s largest refugee camp. Ecological Engineering. https://doi.org/10.1016/j.ecoleng.2025.107612",
    "P017. Liu, J., Zevenbergen, C., Lu, J., Qi, Q., Veerbeek, W., Chowdhury, S. W., & Qian, L. (2026). Land-Use Governance of Borderland Protected Areas Under Refugee Expansion and Climate Threats: Evidence from Teknaf, Bangladesh. Land. https://doi.org/10.3390/land15061024",
    "P018. Sobczak-Szelc, K., Chułek, M., Espegren, A., Malak, M. A., Jenerowicz-Sanikowska, M., & Quader, M. A. (2026). Exploring the Role of Common Pool Resources and (Mal)Coping in Socio-Environmental Dynamics: A Case Study of the Kutupalong Refugee Camp, Bangladesh. Journal of Immigrant & Refugee Studies. https://doi.org/10.1080/15562948.2025.2607538",
    "P019. Rashid, K. J., Hoque, M. A., Esha, T. A., Rahman, M. A., & Paul, A. (2021). Spatiotemporal changes of vegetation and land surface temperature in the refugee camps and its surrounding areas of Bangladesh after the Rohingya influx from Myanmar. Environment, Development and Sustainability. https://doi.org/10.1007/s10668-020-00733-x",
    "Dataset. Hansen, M. C., Potapov, P. V., Moore, R., et al. (2013). High-Resolution Global Maps of 21st-Century Forest Change. Science, 342(6160), 850–853. https://doi.org/10.1126/science.1244693. Product used: Global Forest Change v1.12 (2000–2024).",
    "Dataset. European Commission Joint Research Centre. (2025). Tropical Moist Forest annual change collection (1990–2025), v1_2025. Supporting reference: Vancutsem, C., et al. (2021). Long-term monitoring of tropical moist forest transitions from 1990 to 2019. Science Advances, 7(10), eabe1603. https://doi.org/10.1126/sciadv.abe1603",
    "Dataset. Santoro, M., & Cartus, O. (2025). ESA Biomass Climate Change Initiative (Biomass_cci): Global datasets of forest above-ground biomass for 2005–2012 and 2015–2024, v7.0. CEDA. https://doi.org/10.5285/6429d1aafe1e43b9b414e4a5a7f8b903",
    "Dataset. Dubayah, R. O., Armston, J., Healey, S. P., Yang, Z., Patterson, P. L., Saarela, S., Stahl, G., Duncanson, L., Kellner, J. R., Bruening, J. M., & Pascual, A. (2023). GEDI L4B Gridded Aboveground Biomass Density, Version 2.1. ORNL DAAC. https://doi.org/10.3334/ORNLDAAC/2299",
]
replace_between(doc, references, appendices, [(reference_intro, "Normal")] + [(ref, "Normal") for ref in references_list])

doc.save(PATH)
print("v02 QA corrections applied.")
