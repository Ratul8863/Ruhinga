# -*- coding: utf-8 -*-
"""Apply final boundary-stage consistency fixes and freeze markers in report v05."""
import docx
from copy import deepcopy
from docx.oxml.ns import qn
from docx.shared import Pt

PATH = "01_Report/Full_Research_Report_v05_Boundary_Stage_Frozen.docx"


def find(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise KeyError(text)


def insert_before(anchor, text, style="Normal"):
    p = anchor._parent.add_paragraph()
    p.style = style
    p.add_run(text)
    anchor._p.addprevious(p._p)
    return p


def insert_after(anchor, text, style="Normal"):
    p = anchor._parent.add_paragraph()
    p.style = style
    p.add_run(text)
    anchor._p.addnext(p._p)
    return p


def set_run_size(paragraph, size=9):
    for run in paragraph.runs:
        run.font.size = Pt(size)


doc = docx.Document(PATH)

# Title page
for p in doc.paragraphs:
    if p.text.startswith("Version:"):
        p.text = "Version: v05"
    elif p.text.startswith("Date:"):
        p.text = "Date: 25 July 2026"
    elif p.text.startswith("Stage:"):
        p.text = (
            "Stage: Boundary Verification and Historical Crosswalk Frozen; "
            "Next — Annual Camp Exposure Reconstruction"
        )

# Project status box after stage line
stage = next(p for p in doc.paragraphs if p.text.startswith("Stage:"))
status_lines = [
    "Current Project Status",
    "✓ Literature Review — Completed",
    "✓ Official Boundary Verification — Completed (FROZEN)",
    "✓ Historical Identifier Crosswalk — Completed (FROZEN)",
    "□ Annual Camp Exposure Reconstruction — Pending (Step 3.3)",
    "□ Forest Extraction — Pending",
    "□ Fragmentation — Pending",
    "□ SDID — Pending",
    "□ Biomass / Carbon — Pending",
    "□ Machine Learning — Pending",
    "□ Conservation Priority — Pending",
]
anchor = stage
for line in reversed(status_lines):
    insert_after(anchor, line)

# Restructure 4.4 headings
h44 = find(doc, "4.4 Camp-Exposure Reconstruction")
h44.text = "4.4 Official Camp Boundary Processing"
h441 = find(doc, "4.4.1 Official Boundary Processing and Standardisation")
h441.text = "4.4.1 Geometry Standardisation"
h442 = find(doc, "4.4.2 Harmonisation of the 2018 and 2023 Camp Identifiers")
h442.text = "4.4.2 Historical Identifier Crosswalk"

# Add 4.4.3 before 4.5
h45 = find(doc, "4.5 Forest-Loss Analysis")
insert_before(h45, "4.4.3 Annual Exposure Reconstruction (Future Work)", "Heading 3")
insert_before(
    h45,
    "Purpose — Reconstruct physical annual camp footprints for years without official administrative "
    "boundary products, especially 2016–2017 and other unavailable intermediate years.",
)
insert_before(
    h45,
    "Input Dataset(s) — Locked official 2018, 2023 and 2024 administrative extents; historical "
    "crosswalk (camp_crosswalk_2018_2023_final.csv); satellite imagery collections.",
)
insert_before(
    h45,
    "Method / Tool — Satellite-derived settlement/footprint extraction using locked official layers "
    "as temporal anchors and the frozen crosswalk for identifier consistency (EPSG:32646).",
)
insert_before(
    h45,
    "Output Dataset — Pending annual exposure layers for 2016–2024/2025 under Step 3.3.",
)
insert_before(
    h45,
    "Results / QA Metrics — Not yet generated; this subsection marks planned future work only.",
)
insert_before(
    h45,
    "Interpretation — Official boundary products remain administrative temporal anchors. Physical "
    "camp expansion and exposure effects will not be inferred until annual reconstructed footprints "
    "are completed and validated.",
)

# Strengthen 5.3 nearly-identical sentence
for p in doc.paragraphs:
    if "nearly identical" in p.text and "0.057033" in p.text:
        p.text = (
            "Camp-wise joining between the 2023 A1 and 2024 A2-derived products successfully matched "
            "all 33 CampSSID values, with zero unmatched camps. The signed net area difference was "
            "approximately 0.000027 ha. The summed absolute camp-wise area difference was only "
            "0.057033 ha across all 33 camps, with a maximum individual difference of 0.015016 ha. "
            "These numerical results indicate very close administrative area agreement; however, "
            "area agreement alone is not treated as proof of exact geometric equivalence or physical "
            "settlement stability."
        )
        break

# Add important sentence to 4.12.2 interpretation
for p in doc.paragraphs:
    if p.text.startswith("Interpretation — Shamlapur was the only 2018 site"):
        if "validated crosswalk will be used throughout" not in p.text:
            p.text = (
                p.text
                + " The validated crosswalk will be used throughout all subsequent annual exposure "
                "reconstruction, spatial overlay and causal analyses to ensure temporal consistency "
                "of camp identifiers."
            )
        break

# Replace 5.7 text
h57 = find(doc, "5.7 Empirical Work Remaining")
# next paragraph after heading
ps = list(doc.paragraphs)
idx = next(i for i, p in enumerate(ps) if p._p is h57._p)
if idx + 1 < len(ps):
    ps[idx + 1].text = (
        "The 2018–2023 historical identifier crosswalk has been completed and validated. Remaining "
        "work includes reconstruction of missing annual camp footprints (2016–2017 and other "
        "unavailable years), annual exposure reconstruction, Hansen outcome extraction, "
        "fragmentation analysis, SDID implementation, biomass/carbon estimation, machine-learning "
        "modelling and conservation-priority mapping."
    )

# Rebuild Table B4 with Interpretation column
# Find the last table (B4) and replace content
b4_caption = find(doc, "Table B4. 2018–2023 crosswalk validation")
# The table immediately after caption is tables[-1] in current structure, but safer to locate by position
tbl_element = b4_caption._p.getnext()
while tbl_element is not None and tbl_element.tag != qn("w:tbl"):
    tbl_element = tbl_element.getnext()
if tbl_element is None:
    raise RuntimeError("Table B4 not found after caption.")

# Remove old table
parent = tbl_element.getparent()
parent.remove(tbl_element)

# Insert new table after caption
new_table = doc.add_table(rows=1, cols=3)
new_table.style = "Table Grid"
headers = ["Metric", "Result", "Interpretation"]
for i, h in enumerate(headers):
    new_table.rows[0].cells[i].text = h
rows = [
    ["2018 historical sites", "38", "Complete historical site universe"],
    ["Sites with spatial overlap", "37", "Intersected ≥1 2023 camp polygon"],
    ["Unmatched sites", "1", "Retained explicitly; not force-matched"],
    ["Primary-match ties", "0", "Unique dominant overlap for every match"],
    ["High-confidence matches", "33", "Dominant overlap ≥95%"],
    ["Moderate-confidence matches", "4", "Dominant overlap ≥80% and <95%"],
    ["Low-confidence matches", "0", "No match below 80% overlap"],
    ["Mean primary overlap", "97.70%", "Average dominant-overlap strength"],
    ["Median primary overlap", "99.17%", "Central tendency of match quality"],
    ["Minimum primary overlap", "80.19%", "Weakest accepted Moderate match"],
    ["Maximum primary overlap", "100%", "Perfect spatial coincidence cases"],
    ["Unmatched site", "Shamlapur", "No intersection with 2023 A1 layer"],
    ["Nearest 2023 camp to Shamlapur", "CXB-085 / Camp 22", "Nearest-neighbour diagnostic only"],
    ["Nearest distance", "4,462.79 m", "Too distant for forced assignment"],
]
for values in rows:
    cells = new_table.add_row().cells
    for i, value in enumerate(values):
        cells[i].text = value
for row in new_table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)

# Move newly created table to sit after caption
new_tbl_element = new_table._tbl
# python-docx appends table at document end; relocate after caption
new_tbl_element.getparent().remove(new_tbl_element)
b4_caption._p.addnext(new_tbl_element)

# Freeze note at end of Appendix B
doc.add_paragraph(
    "Boundary-stage freeze note — The verified 2018, 2023 and 2024 official administrative layers, "
    "the 2018–2023 historical crosswalk, associated QA evidence, Decision Log entries and metadata "
    "records are locked. Any future change must be introduced through a new report version "
    "(v06+) without overwriting the validated workflow."
)

doc.save(PATH)
print("v05 freeze report updated.")
