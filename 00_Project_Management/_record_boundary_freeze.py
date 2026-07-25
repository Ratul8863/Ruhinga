# -*- coding: utf-8 -*-
"""Record Decision 036, Progress Log freeze, and Used_in_Analysis metadata."""
from datetime import date

import docx
import openpyxl

TODAY = "25 July 2026"

# --- Metadata: Used_in_Analysis ---
meta_path = "04_Data/02_Camp_Exposure/Metadata/official_boundary_layer_verification.csv"
rows = open(meta_path, encoding="utf-8").read().splitlines()
header = rows[0].split(",")
if "Used_in_Analysis" not in header:
    # Simple CSV with quoted fields — rewrite carefully with csv module
    import csv
    from io import StringIO

    with open(meta_path, encoding="utf-8", newline="") as f:
        reader = list(csv.DictReader(f))
    used_map = {
        "OB01": "Yes",
        "OB02": "Yes",
        "OB03-A2": "Yes",
        "OB03-A3": "Validation only",
        "CW01": "Yes",
    }
    fieldnames = list(reader[0].keys()) + ["Used_in_Analysis"]
    with open(meta_path, "w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in reader:
            row["Used_in_Analysis"] = used_map.get(row["layer_id"], "")
            writer.writerow(row)
    print("Updated metadata Used_in_Analysis")
else:
    print("Used_in_Analysis already present")

# --- Decision 036 ---
dlog = docx.Document("00_Project_Management/Research_Decision_Log.docx")
existing = "\n".join(p.text for p in dlog.paragraphs)
if "Decision 036" not in existing:
    dlog.add_paragraph("Decision 036 - Historical Crosswalk Locked")
    dlog.add_paragraph("Decision:")
    dlog.add_paragraph(
        "The verified 2018–2023 historical identifier crosswalk is locked. "
        "No future manual reassignment of Site–CampSSID matches is permitted "
        "unless official evidence becomes available."
    )
    dlog.add_paragraph("Implementation:")
    dlog.add_paragraph(
        "1. Treat camp_crosswalk_2018_2023_final.csv as the frozen master crosswalk."
    )
    dlog.add_paragraph(
        "2. Use the locked crosswalk for all subsequent annual exposure reconstruction, "
        "spatial overlay and causal analyses."
    )
    dlog.add_paragraph(
        "3. If official evidence later requires change, introduce a new versioned "
        "crosswalk and report revision (v06+) without overwriting the validated workflow."
    )
    dlog.add_paragraph("Reason:")
    dlog.add_paragraph(
        "Thirty-seven of 38 sites have unique primary matches with high/moderate "
        "confidence (High = 33; Moderate = 4; Low = 0). Shamlapur is retained as "
        "explicitly unmatched. Boundary verification, crosswalk QA and documentation "
        "are now complete at journal-ready stage."
    )
    dlog.add_paragraph("Freeze scope:")
    dlog.add_paragraph(
        "2018 official layer LOCK; 2023 official layer LOCK; 2024 A2 primary LOCK; "
        "2024 A3 validation LOCK; crosswalk LOCK; QA LOCK; Decision Log LOCK for "
        "boundary stage; metadata LOCK; Full_Research_Report_v05_Boundary_Stage_Frozen.docx LOCK."
    )
    dlog.save("00_Project_Management/Research_Decision_Log.docx")
    print("Added Decision 036")
else:
    print("Decision 036 already present")

# --- Progress Log ---
wb = openpyxl.load_workbook("00_Project_Management/Research_Progress_Log.xlsx")
ws = wb["Progress_Log"]

# Update the outdated Pending row for Step 3.2B if still Pending
for row in ws.iter_rows(min_row=2):
    step = row[1].value or ""
    task = row[2].value or ""
    status = row[5].value or ""
    if (
        step == "Step 3.2B"
        and "Next task identified" in str(task)
        and status == "Pending"
    ):
        row[3].value = (
            "Superseded: 2018–2023 crosswalk was completed and validated under "
            "subsequent Step 3.2B tasks; boundary stage frozen."
        )
        row[4].value = "camp_crosswalk_2018_2023_final.csv; Research_Decision_Log.docx (034–036)"
        row[5].value = "Superseded"
        row[6].value = "Earlier pending note retained for audit trail only."
        row[7].value = "Proceed to Step 3.3 — Annual Camp Exposure Reconstruction."

# Append freeze + next-step rows if not already present
existing_tasks = {
    (str(r[1] or ""), str(r[2] or ""))
    for r in ws.iter_rows(min_row=2, values_only=True)
}
new_rows = [
    (
        TODAY,
        "Step 3.2B",
        "Boundary stage frozen after crosswalk QA",
        "Locked verified 2018/2023/2024 official layers, 2018–2023 crosswalk, QA evidence, "
        "Decision 036 and report v05; no overwrite of validated boundary workflow.",
        "Full_Research_Report_v05_Boundary_Stage_Frozen.docx; "
        "official_boundary_layer_verification.csv; Research_Decision_Log.docx (Decision 036)",
        "Completed",
        "None — boundary-processing phase complete.",
        "Begin Step 3.3 — Annual Camp Exposure Reconstruction (2016–2025 exposure timeline).",
    ),
    (
        TODAY,
        "Step 3.3",
        "Annual Camp Exposure Reconstruction",
        "Not started. Next major empirical step after frozen boundary stage.",
        "Pending",
        "Pending",
        "Missing annual footprints for 2016–2017 and other unavailable years.",
        "Reconstruct annual camp footprints using locked official anchors and frozen crosswalk; start with 2016 pre-influx satellite baseline.",
    ),
]
for row in new_rows:
    key = (row[1], row[2])
    if key not in existing_tasks:
        ws.append(list(row))
        print("Appended:", row[1], "-", row[2])
    else:
        print("Already present:", row[1], "-", row[2])

wb.save("00_Project_Management/Research_Progress_Log.xlsx")
print("Progress Log updated")
