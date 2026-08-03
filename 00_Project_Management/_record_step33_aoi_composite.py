# -*- coding: utf-8 -*-
"""Record Decisions 037–039, Progress Log 3.3.1–3.3.2, and reconstruction metadata."""
import csv
from pathlib import Path

import docx
import openpyxl

TODAY = "27 July 2026"
ROOT = Path(__file__).resolve().parents[1]

# ---------------------------------------------------------------------------
# Decision Log 037–039
# ---------------------------------------------------------------------------
dlog_path = ROOT / "00_Project_Management" / "Research_Decision_Log.docx"
dlog = docx.Document(str(dlog_path))
existing = "\n".join(p.text for p in dlog.paragraphs)

entries = [
    (
        "Decision 037 - Reconstruction Processing AOI",
        [
            ("Decision:", None),
            (
                None,
                "The annual satellite reconstruction processing AOI will use the union of "
                "the locked 2018, 2023 and 2024 A2-derived official camp extents plus a "
                "2,000 m dissolved buffer.",
            ),
            (
                None,
                "The AOI is a processing extent only and is not a causal treatment, "
                "spillover or donor-zone definition.",
            ),
            ("Implementation:", None),
            (
                None,
                "1. Merge the locked 2018, 2023 and 2024 A2-derived official camp layers "
                "(104 source features).",
            ),
            (
                None,
                "2. Dissolve without a grouping field into a single multipart geometry.",
            ),
            (
                None,
                "3. Buffer by 2,000 m in EPSG:32646 using round joins, 20 segments and a "
                "dissolved output.",
            ),
            (
                None,
                "4. Export reconstruction_aoi_2016_2024_2km_final.gpkg and upload the "
                "Earth Engine asset reconstruction_aoi_2016_2024_2km.",
            ),
            ("Validation:", None),
            (
                None,
                "One valid multipart feature; 0 invalid; 0 geometry errors; area "
                "approximately 189.372 km².",
            ),
            ("Reason:", None),
            (
                None,
                "A shared processing AOI is required for consistent annual satellite "
                "composites, while preserving the locked official layers as administrative "
                "anchors rather than physical footprints.",
            ),
        ],
    ),
    (
        "Decision 038 - 2016 Baseline Window",
        [
            ("Decision:", None),
            (
                None,
                "The 2016 pre-influx baseline composite will use imagery from "
                "1 November 2015 to 30 April 2016.",
            ),
            ("Reason:", None),
            (
                None,
                "The 42-scene window achieved 99.738% valid-pixel coverage and the lowest "
                "mean scene cloud cover among the three candidate windows, while remaining "
                "within a dry-season period that excludes post-influx dates.",
            ),
            ("Condition:", None),
            (
                None,
                "The product is documented as a 2016 pre-influx settlement and land-surface "
                "baseline, not as a 2016 Rohingya camp footprint, until settlement "
                "classification and physical-footprint delineation are completed.",
            ),
        ],
    ),
    (
        "Decision 039 - 2016 Composite Specification",
        [
            ("Decision:", None),
            (
                None,
                "The baseline composite will use Landsat 8 Collection 2 Tier 1 Level 2 "
                "Surface Reflectance, QA-based cloud masking, a median composite, 30 m "
                "output resolution and EPSG:32646.",
            ),
            (
                None,
                "The output will retain six reflectance bands, NDVI, NDBI, MNDWI, BSI and "
                "valid observation count.",
            ),
            ("Implementation:", None),
            (
                None,
                "Export L8_SR_2016_preinflux_windowB_composite_30m.tif to "
                "04_Data/02_Camp_Exposure/Annual_Reconstruction/03_Annual_Composites/ "
                "using GEE scripts 01_2016_Landsat_Window_Test and "
                "02_2016_Landsat_Composite_Export.",
            ),
            ("Validation:", None),
            (
                None,
                "11 Float32 bands; 610 × 1,259 pixels; EPSG:32646; 30 m; valid observations "
                "13–31 (mean ≈ 26.02); AOI valid-pixel coverage ≈ 99.738%.",
            ),
        ],
    ),
]

if "Decision 037" not in existing:
    for title, parts in entries:
        dlog.add_paragraph(title)
        for label, text in parts:
            if label and text is None:
                dlog.add_paragraph(label)
            elif text is not None:
                dlog.add_paragraph(text)
    dlog.save(str(dlog_path))
    print("Added Decisions 037–039")
else:
    print("Decisions 037+ already present")

# ---------------------------------------------------------------------------
# Progress Log
# ---------------------------------------------------------------------------
plog_path = ROOT / "00_Project_Management" / "Research_Progress_Log.xlsx"
wb = openpyxl.load_workbook(str(plog_path))
ws = wb["Progress_Log"]
existing_tasks = {
    (str(r[1] or ""), str(r[2] or ""))
    for r in ws.iter_rows(min_row=2, values_only=True)
}

# Update the earlier pending Step 3.3 row if present
for row in ws.iter_rows(min_row=2):
    step = str(row[1].value or "")
    task = str(row[2].value or "")
    status = str(row[5].value or "")
    if step == "Step 3.3" and task == "Annual Camp Exposure Reconstruction" and status == "Pending":
        row[3].value = (
            "Initiated: reconstruction AOI validated and 2016 pre-influx Landsat baseline "
            "composite exported; settlement classification still pending."
        )
        row[4].value = (
            "reconstruction_aoi_2016_2024_2km_final.gpkg; "
            "L8_SR_2016_preinflux_windowB_composite_30m.tif; "
            "Full_Research_Report_v06_Annual_Exposure_Reconstruction.docx"
        )
        row[5].value = "In Progress"
        row[6].value = (
            "No physical 2016 camp footprint yet; classification and accuracy assessment pending."
        )
        row[7].value = (
            "Proceed to settlement training-data preparation and 2016 physical-footprint "
            "classification (Step 3.3.3)."
        )

new_rows = [
    (
        TODAY,
        "Step 3.3.1",
        "Created and validated 2 km reconstruction AOI",
        "Merged locked 2018/2023/2024 A2 official camp layers (104 features), dissolved to "
        "one multipart geometry, buffered 2,000 m in EPSG:32646; strict validity 1/0/0; "
        "area ≈ 189.372 km². Uploaded EE asset reconstruction_aoi_2016_2024_2km.",
        "04_Data/02_Camp_Exposure/Annual_Reconstruction/01_AOI/"
        "reconstruction_aoi_2016_2024_2km_final.gpkg; "
        "official_camp_snapshots_2018_2024_merged.gpkg; "
        "official_camp_snapshots_2018_2024_union.gpkg; "
        "reconstruction_aoi_2016_2024_2km_wgs84.zip",
        "Complete",
        "AOI is processing extent only — not treatment/spillover/donor zone.",
        "Select 2016 seasonal window and export baseline composite.",
    ),
    (
        TODAY,
        "Step 3.3.2",
        "Selected 2016 window and exported 11-band composite",
        "Compared three Landsat 8 C2 T1 L2 windows; selected Nov 2015–Apr 2016 "
        "(42 scenes; mean cloud 16.96%; AOI valid-pixel coverage 99.738%). Exported "
        "median SR composite with 6 reflectance bands + NDVI/NDBI/MNDWI/BSI + valid count.",
        "04_Data/02_Camp_Exposure/Annual_Reconstruction/03_Annual_Composites/"
        "L8_SR_2016_preinflux_windowB_composite_30m.tif; "
        "GEE scripts 01_2016_Landsat_Window_Test; 02_2016_Landsat_Composite_Export; "
        "Research_Decision_Log.docx (037–039); "
        "Full_Research_Report_v06_Annual_Exposure_Reconstruction.docx",
        "Composite QA complete; classification pending",
        "Documented as 2016 pre-influx settlement/land-surface baseline — not a camp footprint.",
        "Prepare training data and classify 2016 settlement / physical footprint.",
    ),
]

for row in new_rows:
    key = (row[1], row[2])
    if key not in existing_tasks:
        ws.append(list(row))
        print("Appended Progress:", row[1], "-", row[2])
    else:
        print("Progress already present:", row[1], "-", row[2])

wb.save(str(plog_path))
print("Progress Log saved")

# ---------------------------------------------------------------------------
# Metadata inventory for Annual Reconstruction
# ---------------------------------------------------------------------------
meta_dir = ROOT / "04_Data" / "02_Camp_Exposure" / "Annual_Reconstruction" / "07_Metadata"
meta_dir.mkdir(parents=True, exist_ok=True)
meta_path = meta_dir / "annual_reconstruction_layer_inventory.csv"

fieldnames = [
    "ID",
    "Dataset",
    "Role",
    "Status",
    "Path",
    "CRS",
    "Notes",
    "Decision",
]

records = [
    {
        "ID": "RAOI01",
        "Dataset": "reconstruction_aoi_2016_2024_2km_final.gpkg",
        "Role": "Satellite processing AOI",
        "Status": "Validated",
        "Path": "04_Data/02_Camp_Exposure/Annual_Reconstruction/01_AOI/",
        "CRS": "EPSG:32646",
        "Notes": (
            "Union of locked 2018/2023/2024 A2 extents + 2 km dissolved buffer; "
            "1 multipart; 189.372 km²; EE asset "
            "projects/rohingya-forest-impact-2026/assets/reconstruction_aoi_2016_2024_2km. "
            "Not treatment/spillover/donor zone."
        ),
        "Decision": "037",
    },
    {
        "ID": "IMG2016-01",
        "Dataset": "L8_SR_2016_preinflux_windowB_composite_30m.tif",
        "Role": "2016 pre-influx baseline composite",
        "Status": "Raster QA passed",
        "Path": "04_Data/02_Camp_Exposure/Annual_Reconstruction/03_Annual_Composites/",
        "CRS": "EPSG:32646",
        "Notes": (
            "Landsat 8 C2 T1 L2 SR median composite; Nov 2015–Apr 2016; 42 scenes; "
            "11 Float32 bands; 30 m; 610×1259; AOI valid-pixel coverage 99.738%. "
            "Not a 2016 Rohingya camp footprint."
        ),
        "Decision": "038; 039",
    },
]

existing_ids = set()
if meta_path.exists():
    with meta_path.open(encoding="utf-8", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            existing_ids.add(row.get("ID"))

mode = "a" if meta_path.exists() else "w"
with meta_path.open(mode, encoding="utf-8", newline="") as f:
    writer = csv.DictWriter(f, fieldnames=fieldnames)
    if mode == "w":
        writer.writeheader()
    for rec in records:
        if rec["ID"] in existing_ids:
            print("Metadata already present:", rec["ID"])
            continue
        writer.writerow(rec)
        print("Wrote metadata:", rec["ID"])

print("Done.")
